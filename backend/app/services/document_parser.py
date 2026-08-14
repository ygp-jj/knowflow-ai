"""文档解析服务：将上传文件转换为带页码的纯文本页列表。

   这个模块的职责非常单一：把用户上传的【原始文件】转换成【纯文本 + 页码结构】。
   它是整个 RAG 流程的“第一道门”，后面接的是文本切片（split_text）。
   前端理解了这个，就知道为什么上传 PDF 能看到分页，而上传 Word 只有一整块。
"""

from pathlib import Path
from typing import Iterator

from docx import Document as DocxDocument
from docx.oxml.ns import qn
from docx.table import Table
from docx.text.paragraph import Paragraph
from pypdf import PdfReader


class UnsupportedDocumentTypeError(ValueError):
    """不支持的文档类型异常。

    当前端上传了 `.ppt`、`.xlsx` 或图片等格式时，后端会抛出此异常。
    前端捕获到该错误时，建议直接弹窗提示用户：
    “暂不支持解析该文件类型，请上传 PDF、Word 或 TXT 格式。”
    """


def parse_document(file_path: str) -> list[dict]:
    """根据文件扩展名分发到不同的解析器。

    这是模块的入口函数，前端不需要直接调用（由后端服务调用），
    但你需要知道它支持的格式和返回结构：

    支持格式：
        - PDF (.pdf)：逐页提取文字
        - Word (.docx / .doc)：按文档顺序提取段落与表格，【合并为一页】
        - 纯文本 (.md / .txt / .markdown)：读取全文，【合并为一页】

    返回数据结构：
        [
            {
                "page_number": 1,        # 页码（整数），PDF 逐页递增
                "content": "这是该页的文字内容..."
            },
            ...
        ]

    注意：
        - Word 和 TXT 没有“分页”概念，统一返回 page_number = None
        - 如果上传的是不支持的类型，会抛出 UnsupportedDocumentTypeError
    """
    suffix = Path(file_path).suffix.lower()

    if suffix == ".pdf":
        return parse_pdf(file_path)

    if suffix in {".docx", ".doc"}:
        return parse_docx(file_path)

    if suffix in {".md", ".txt", ".markdown"}:
        return parse_text(file_path)

    raise UnsupportedDocumentTypeError(f"暂不支持解析的文件类型: {suffix or '(无扩展名)'}")


def parse_pdf(file_path: str) -> list[dict]:
    """解析 PDF，按页提取文本。

    对 PDF 的处理逻辑：
        1. 按页循环读取
        2. 每页提取的文本放到单独的字典中
        3. 页码从 1 开始（与用户肉眼看到的页码一致）

    前端展示建议：
        可用于实现“分页预览”功能，让用户看到哪段文字来自哪一页。
    """
    reader = PdfReader(file_path)
    pages = []

    for index, page in enumerate(reader.pages):
        text = page.extract_text() or ""   # 某些扫描件可能提取不到文字
        pages.append({
            "page_number": index + 1,       # 页码从 1 开始
            "content": text,
        })

    return pages


def _iter_docx_blocks(document: DocxDocument) -> Iterator[Paragraph | Table]:
    """按 Word 正文顺序产出段落与表格。

    python-docx 的 ``document.paragraphs`` 不含表格内文字；
    必须遍历 body 子节点，才能保留「标题 → 表格 → 下一段」的阅读顺序。
    """
    body = document.element.body
    for child in body.iterchildren():
        if child.tag == qn("w:p"):
            yield Paragraph(child, document)
        elif child.tag == qn("w:tbl"):
            yield Table(child, document)


def _cell_text(cell) -> str:
    """提取单元格纯文本；多段用空格拼接，去掉首尾空白。"""
    parts = [paragraph.text.strip() for paragraph in cell.paragraphs if paragraph.text.strip()]
    return " ".join(parts).replace("\n", " ").strip()


def _table_to_text(table: Table) -> str:
    """将 Word 表格转为可读纯文本（列用 | 分隔，行用换行分隔）。

    合并单元格时 python-docx 会对同一 ``w:tc`` 重复返回，这里按底层节点去重，
    避免同一格文字在一行中出现多次。
    """
    rows: list[str] = []
    for row in table.rows:
        cells: list[str] = []
        seen_tc: set[int] = set()
        for cell in row.cells:
            tc_id = id(cell._tc)
            if tc_id in seen_tc:
                continue
            seen_tc.add(tc_id)
            cells.append(_cell_text(cell))
        if any(cells):
            rows.append(" | ".join(cells))
    return "\n".join(rows)


def parse_docx(file_path: str) -> list[dict]:
    """解析 Word 文档：按文档顺序拼接段落与表格文本。

    对 Word 的处理逻辑：
        1. 按正文顺序遍历段落与表格（不再只用 document.paragraphs）
        2. 表格转为「列用 |、行用换行」的纯文本，避免审批表等结构化内容丢失
        3. 过滤空段落后用换行符拼接
        4. 整个文档合并为【一段】文本（无分页概念）

    前端展示建议：
        由于 Word 没有明确的页码边界，后续切片（split_text）会按字符数切分。
        前端如果展示来源，一般显示为“全文”而不是“第 X 页”。
    """
    document = DocxDocument(file_path)
    parts: list[str] = []
    for block in _iter_docx_blocks(document):
        if isinstance(block, Paragraph):
            text = block.text.strip()
            if text:
                parts.append(text)
            continue

        table_text = _table_to_text(block)
        if table_text:
            parts.append(table_text)

    text = "\n".join(parts)
    return [{"page_number": None, "content": text}]   # 页码置空，表示没有分页信息


def parse_text(file_path: str) -> list[dict]:
    """解析 TXT / Markdown 纯文本文件。

    处理逻辑：
        1. 以 UTF-8 编码读取文件全部内容
        2. 整份文件合并为【一段】文本

    注意：
        - 如果文件不是 UTF-8 编码，这里可能报错（需要前端/用户确认文件编码）。
        - 同样没有页码概念，page_number = None。
    """
    text = Path(file_path).read_text(encoding="utf-8")
    return [{"page_number": None, "content": text}]