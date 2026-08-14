"""文档解析服务单测。"""

import tempfile
import unittest
from pathlib import Path

from docx import Document as DocxDocument

from app.services.document_parser import UnsupportedDocumentTypeError, parse_document, parse_text
from app.services.text_splitter import split_pages_to_chunks


class DocumentParserTests(unittest.TestCase):
    def test_parse_text_and_markdown(self):
        with tempfile.TemporaryDirectory() as temp_dir:
            txt_path = Path(temp_dir) / "note.txt"
            md_path = Path(temp_dir) / "note.md"
            txt_path.write_text("hello txt", encoding="utf-8")
            md_path.write_text("# hello md", encoding="utf-8")

            self.assertEqual(parse_text(str(txt_path)), [{"page_number": None, "content": "hello txt"}])
            self.assertEqual(
                parse_document(str(md_path)),
                [{"page_number": None, "content": "# hello md"}],
            )

    def test_parse_docx(self):
        with tempfile.TemporaryDirectory() as temp_dir:
            docx_path = Path(temp_dir) / "demo.docx"
            document = DocxDocument()
            document.add_paragraph("第一段")
            document.add_paragraph("第二段")
            document.save(docx_path)

            pages = parse_document(str(docx_path))
            self.assertEqual(len(pages), 1)
            self.assertIsNone(pages[0]["page_number"])
            self.assertIn("第一段", pages[0]["content"])
            self.assertIn("第二段", pages[0]["content"])

    def test_parse_docx_keeps_table_in_document_order(self):
        """Word 正文中的表格必须进入解析结果，且夹在前后段落之间。"""
        with tempfile.TemporaryDirectory() as temp_dir:
            docx_path = Path(temp_dir) / "leave-policy.docx"
            document = DocxDocument()
            document.add_paragraph("第三章 请假审批权限")
            document.add_paragraph("第十三条 审批权限")
            table = document.add_table(rows=3, cols=2)
            table.rows[0].cells[0].text = "请假天数"
            table.rows[0].cells[1].text = "审批流程"
            table.rows[1].cells[0].text = "1天以内"
            table.rows[1].cells[1].text = "直属主管审批"
            table.rows[2].cells[0].text = "2-3天"
            table.rows[2].cells[1].text = "部门负责人审批"
            document.add_paragraph("第十四条 连续请假按累计天数计算审批权限")
            document.save(docx_path)

            pages = parse_document(str(docx_path))
            content = pages[0]["content"]

            self.assertIn("第十三条 审批权限", content)
            self.assertIn("请假天数 | 审批流程", content)
            self.assertIn("1天以内 | 直属主管审批", content)
            self.assertIn("2-3天 | 部门负责人审批", content)
            self.assertIn("第十四条 连续请假按累计天数计算审批权限", content)
            self.assertLess(content.index("第十三条"), content.index("请假天数"))
            self.assertLess(content.index("部门负责人审批"), content.index("第十四条"))

            chunks = split_pages_to_chunks(pages, chunk_size=256, chunk_overlap=50)
            joined = "\n".join(item["content"] for item in chunks)
            self.assertIn("直属主管审批", joined)
            # 大标题独立成块；表格作为子块挂在第十三条下
            article_chunks = [item for item in chunks if item["content"].strip() == "第十三条 审批权限"]
            self.assertEqual(len(article_chunks), 1)
            table_chunks = [item for item in chunks if "直属主管审批" in item["content"]]
            self.assertEqual(len(table_chunks), 1)
            self.assertEqual(
                table_chunks[0].get("parent_chunk_index"),
                article_chunks[0]["chunk_index"],
            )

    def test_unsupported_xlsx_raises(self):
        with tempfile.TemporaryDirectory() as temp_dir:
            xlsx_path = Path(temp_dir) / "sheet.xlsx"
            xlsx_path.write_bytes(b"not-a-real-xlsx")

            with self.assertRaises(UnsupportedDocumentTypeError) as ctx:
                parse_document(str(xlsx_path))

            self.assertIn("暂不支持", str(ctx.exception))


if __name__ == "__main__":
    unittest.main()
